from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import streamlit as st
import io
from packages.utils import get_folder_id, get_image_map, download_image_as_bytes

def set_columns(section, cols):
    """
    Sets the number of columns for a section using OXML.
    """
    sectPr = section._sectPr
    cols_xml = sectPr.xpath('./w:cols')[0] if sectPr.xpath('./w:cols') else None
    
    if cols_xml is None:
        cols_xml = OxmlElement('w:cols')
        sectPr.append(cols_xml)
        
    cols_xml.set(qn('w:num'), str(cols))
    cols_xml.set(qn('w:space'), "720") # Space between columns (720 twips = 0.5 inch = 1.27cm)

def add_header(section, info_text):
    """
    Add header with dynamic info and student fields.
    """
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.clear()
    
    # Tab stops for alignment
    # B4 landscape width = 36.4cm. Margins = 1.27cm on each side.
    # Content width = 33.86cm.
    # Center tab ~ 17cm. Right tab ~ 33.8cm.
    tab_stops = paragraph.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(17), WD_TAB_ALIGNMENT.CENTER)
    tab_stops.add_tab_stop(Cm(33.8), WD_TAB_ALIGNMENT.RIGHT)
    
    # Left: Info Text
    run_left = paragraph.add_run(info_text)
    run_left.font.name = "微軟正黑體"
    run_left.font.size = Pt(10)
    
    # Right: Student Fields
    run_right = paragraph.add_run("\t\t班級：_______ 座號：_______ 姓名：_________________")
    run_right.font.name = "微軟正黑體"
    run_right.font.size = Pt(10)

def add_footer(section):
    """
    Add footer with page numbers.
    """
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def add_questions(doc, selected_data, image_map, mode="student", total_questions=0, start_idx=0, progress_bar=None):
    """
    Add questions to the document.
    mode: "student" (hide answer) or "teacher" (show answer)
    """
    
    for i, (_, row) in enumerate(selected_data.iterrows()):
        # Update progress if provided
        if progress_bar:
            # We have 2 phases, so scale progress
            current = start_idx + i + 1
            progress_val = min(current / total_questions, 1.0)
            progress_bar.progress(progress_val)

        # --- Question Block ---
        
        # 1. Text Info (Moved above image)
        # Format: 【Year Source】 No. Unit Difficulty
        difficulties = row.get('難易度', '')
        info_parts = [
            f"【{row.get('年份')} {row.get('來源')}】",
            f"題號：{row.get('題號')}",
            f"單元：{row.get('單元')}",
            f"難易度：{difficulties}"
        ]
        
        # Add Answer/Rate if Teacher mode
        if mode == "teacher":
            info_parts.append(f"答案：{row.get('答案')}")
            info_parts.append(f"答對率：{row.get('答對率', 'N/A')}")
            
        info_text = "  ".join(info_parts)
        
        p_info = doc.add_paragraph()
        p_info.keep_with_next = True # Keep info with image
        runner = p_info.add_run(info_text)
        runner.font.size = Pt(10)
        runner.font.name = "微軟正黑體"
        
        # 2. Image
        image_name = row.get("圖檔名")
        image_fh = None
        
        if image_name in image_map:
            image_fh = download_image_as_bytes(image_map[image_name])
            
        p_img = doc.add_paragraph()
        p_img.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p_img.paragraph_format.space_after = Pt(18) # Space after item
        
        if image_fh:
            try:
                # Resize to fit column width (approx 16cm for B4 2-col)
                p_img.add_run().add_picture(image_fh, width=Cm(16))
            except Exception as e:
                p_img.add_run(f"[圖片載入失敗: {image_name}]")
        else:
            p_img.add_run(f"[找不到圖片: {image_name}]")

def generate_b4_word(selected_data, filter_info_str):
    """
    Generates a B4 Word document with Student/Teacher versions.
    """
    doc = Document()
    
    # 0. Prepare Resources
    folder_id = get_folder_id("Math_Crops")
    if not folder_id:
        st.error("找不到題目圖檔資料夾 'Math_Crops'！")
        return None
        
    with st.spinner("正在索引圖檔資料庫..."):
        image_map = get_image_map(folder_id)

    total_items = len(selected_data) * 2 # 2 passes
    progress_bar = st.progress(0)
    
    # --- Part 1: Student Version ---
    section_student = doc.sections[0]
    
    # 1. Page Setup
    section_student.page_width = Cm(36.4)
    section_student.page_height = Cm(25.7)
    section_student.left_margin = Cm(1.27)
    section_student.right_margin = Cm(1.27)
    section_student.top_margin = Cm(1.27)
    section_student.bottom_margin = Cm(1.27)
    set_columns(section_student, 2)
    
    # 2. Header & Footer
    # Remove header as requested to save space
    # add_header(section_student, f"{filter_info_str} (學用)")
    add_footer(section_student)
    
    # 3. Content
    # Add Info Block directly in the body (Top of Left Column)
    add_info_block(doc, f"{filter_info_str} (學用)")
    
    add_questions(doc, selected_data, image_map, mode="student", 
                 total_questions=total_items, start_idx=0, progress_bar=progress_bar)
    
    # --- Part 2: Teacher Version ---
    doc.add_page_break()
    
    # Add Info Block for Teacher Version
    add_info_block(doc, f"{filter_info_str} (教用)")
    
    add_questions(doc, selected_data, image_map, mode="teacher", 
                 total_questions=total_items, start_idx=len(selected_data), progress_bar=progress_bar)

    # Save
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

def add_info_block(doc, info_text):
    """
    Adds the filter info and student fields at the top of the column.
    replaces the traditional header.
    """
    # Line 1: Filter Info
    p_info = doc.add_paragraph()
    run_info = p_info.add_run(info_text)
    run_info.font.name = "微軟正黑體"
    run_info.font.size = Pt(10)
    run_info.bold = True
    p_info.paragraph_format.space_after = Pt(2) # minimal space
    
    # Line 2: Student Fields
    p_fields = doc.add_paragraph()
    run_fields = p_fields.add_run("班級：_______ 座號：_______ 姓名：_________________")
    run_fields.font.name = "微軟正黑體"
    run_fields.font.size = Pt(10)
    p_fields.paragraph_format.space_after = Pt(12) # Space before first question
    
    # Add a separator line (optional, but good for distinction)
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_after = Pt(12)
    border = p_line.add_run()
    border.add_break() # Just a small break if we don't use real borders
    # Actually, let's just use the space. Removing the explicit line to save space as requested.
    # We'll rely on space_after of p_fields.
    
    # Remove p_line from doc (undo add)
    # Since we can't easily undo in python-docx list, just don't add it.
    # p_fields space_after is enough.


def generate_a4_word(selected_data, filter_info_str, include_answer=True):
    """
    Generates an A4 Word document (Single Column).
    Includes Student Version, and optional Teacher Version.
    """
    doc = Document()
    
    # 0. Prepare Resources
    folder_id = get_folder_id("Math_Crops")
    if not folder_id:
        return None
        
    # We might not need spinner if cached, but for safety
    image_map = get_image_map(folder_id)

    total_items = len(selected_data) * (2 if include_answer else 1)
    # progress_bar = st.progress(0) # Optional, or pass None
    
    # --- Part 1: Student Version ---
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    
    add_footer(section)
    
    add_info_block(doc, f"{filter_info_str} (自我練習卷)")
    
    add_questions(doc, selected_data, image_map, mode="student", 
                 total_questions=total_items, start_idx=0, progress_bar=None)
                 
    # --- Part 2: Teacher Version ---
    if include_answer:
        doc.add_page_break()
        add_info_block(doc, f"{filter_info_str} (解析卷)")
        add_questions(doc, selected_data, image_map, mode="teacher", 
                     total_questions=total_items, start_idx=len(selected_data), progress_bar=None)

    # Save
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer
